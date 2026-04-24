from __future__ import annotations

import json
import re
import shutil
import time
import urllib.error
import urllib.parse
import urllib.request
from collections import defaultdict
from dataclasses import dataclass
from datetime import date, datetime, timedelta, timezone
from decimal import Decimal
from pathlib import Path
from zoneinfo import ZoneInfo

from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter


BASE = Path(r"C:\Users\d.gilland\Documents\Codex\2026-04-24\files-mentioned-by-the-user-codex")
DOC = Path(
    r"C:\Users\d.gilland\OneDrive - VPC Pizza Management LLC\Codex\Codex_shopify, klaviyo, meta, gmail, postpilot.docx"
)
OUT = BASE / "giordanos_dtc_daily_revenue_by_channel_2026-03-03_to_2026-04-23.xlsx"

START = date(2026, 3, 3)
END = date(2026, 4, 23)
LOCAL_TZ = ZoneInfo("America/Chicago")
KLAVIYO_REVISION = "2025-10-15"
KLAVIYO_CONVERSION_METRIC_ID = "YadxyL"  # Shopify Placed Order in this Klaviyo account.


@dataclass
class Credentials:
    shopify_store: str
    shopify_token: str
    klaviyo_key: str
    meta_account: str
    meta_token: str


def read_credentials() -> Credentials:
    tmp_doc = None
    try:
        doc = Document(DOC)
    except PermissionError:
        tmp_doc = BASE / "credentials_tmp.docx"
        shutil.copy2(DOC, tmp_doc)
        doc = Document(tmp_doc)
    text_parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            text_parts.extend(cell.text for cell in row.cells)
    text = "\n".join(text_parts)
    if tmp_doc and tmp_doc.exists():
        tmp_doc.unlink()

    return Credentials(
        shopify_store=re.search(r"giordanos-frozen-pizza\.myshopify\.com", text).group(0),
        shopify_token=re.search(r"shpat_[A-Za-z0-9]+", text).group(0),
        klaviyo_key=re.search(r"pk_[A-Za-z0-9]+", text).group(0),
        meta_account=re.search(r"act_\d+", text).group(0),
        meta_token=re.search(r"EAAS[A-Za-z0-9]+", text).group(0),
    )


def request_json(url: str, *, headers=None, params=None, method="GET", body=None, timeout=60):
    if params:
        sep = "&" if "?" in url else "?"
        url = f"{url}{sep}{urllib.parse.urlencode(params)}"
    data = json.dumps(body).encode("utf-8") if body is not None else None
    req_headers = dict(headers or {})
    if body is not None:
        req_headers.setdefault("content-type", "application/json")
    req = urllib.request.Request(url, data=data, headers=req_headers, method=method)
    with urllib.request.urlopen(req, timeout=timeout) as resp:
        content = resp.read().decode("utf-8")
        return json.loads(content) if content else {}


def request_json_with_headers(url: str, *, headers=None, params=None, method="GET", body=None, timeout=60):
    if params:
        sep = "&" if "?" in url else "?"
        url = f"{url}{sep}{urllib.parse.urlencode(params)}"
    data = json.dumps(body).encode("utf-8") if body is not None else None
    req_headers = dict(headers or {})
    if body is not None:
        req_headers.setdefault("content-type", "application/json")
    req = urllib.request.Request(url, data=data, headers=req_headers, method=method)
    with urllib.request.urlopen(req, timeout=timeout) as resp:
        content = resp.read().decode("utf-8")
        return (json.loads(content) if content else {}), resp.headers


def next_link(headers) -> str | None:
    link = headers.get("Link")
    if not link:
        return None
    for part in link.split(","):
        if 'rel="next"' in part:
            match = re.search(r"<([^>]+)>", part)
            return match.group(1) if match else None
    return None


def money(value) -> float:
    if value in (None, ""):
        return 0.0
    return float(Decimal(str(value)))


def all_dates() -> list[date]:
    days = []
    d = START
    while d <= END:
        days.append(d)
        d += timedelta(days=1)
    return days


def klaviyo_headers(creds: Credentials) -> dict[str, str]:
    return {
        "Authorization": f"Klaviyo-API-Key {creds.klaviyo_key}",
        "accept": "application/json",
        "revision": KLAVIYO_REVISION,
    }


def clean_topic(name: str) -> str:
    topic = re.sub(r"^TKB\s*\|\s*", "", name or "", flags=re.I)
    topic = re.sub(r"\[(Email|SMS)\]\s*", "", topic, flags=re.I)
    topic = re.sub(r"\s*\(\d{1,2}/\d{1,2}\)\s*$", "", topic)
    return topic.strip()


def date_from_campaign_name(name: str) -> date | None:
    match = re.search(r"\((\d{1,2})/(\d{1,2})\)\s*$", name or "")
    if not match:
        return None
    month = int(match.group(1))
    day = int(match.group(2))
    return date(2026, month, day)


def fetch_klaviyo_campaigns(creds: Credentials):
    headers = klaviyo_headers(creds)
    body = {
        "data": {
            "type": "campaign-values-report",
            "attributes": {
                "timeframe": {
                    "start": "2026-03-03T00:00:00-06:00",
                    "end": "2026-04-23T23:59:59-05:00",
                },
                "conversion_metric_id": KLAVIYO_CONVERSION_METRIC_ID,
                "statistics": ["recipients", "delivered", "conversion_value"],
                "group_by": [
                    "campaign_message_id",
                    "campaign_id",
                    "campaign_message_name",
                    "send_channel",
                ],
            },
        }
    }
    report = request_json(
        "https://a.klaviyo.com/api/campaign-values-reports/",
        headers=headers,
        method="POST",
        body=body,
    )
    results = report["data"]["attributes"].get("results", [])

    by_day = defaultdict(lambda: defaultdict(lambda: {"names": [], "topics": [], "revenue": 0.0}))
    campaign_dates: dict[str, date | None] = {}

    for item in results:
        group = item.get("groupings", {})
        campaign_id = group.get("campaign_id")
        channel = (group.get("send_channel") or "").lower()
        name = group.get("campaign_message_name") or ""
        stats = item.get("statistics", {})
        if channel not in {"email", "sms"} or not campaign_id:
            continue
        if campaign_id not in campaign_dates:
            parsed_day = date_from_campaign_name(name)
            if parsed_day:
                campaign_dates[campaign_id] = parsed_day
                day = parsed_day
            else:
                day = None
        else:
            day = campaign_dates[campaign_id]
        if day is None and campaign_id not in campaign_dates:
            meta = request_json(f"https://a.klaviyo.com/api/campaigns/{campaign_id}/", headers=headers)
            attrs = meta["data"]["attributes"]
            send_time = attrs.get("send_time") or attrs.get("scheduled_at") or attrs.get("created_at")
            campaign_dates[campaign_id] = (
                datetime.fromisoformat(send_time.replace("Z", "+00:00")).astimezone(LOCAL_TZ).date()
                if send_time
                else None
            )
            time.sleep(0.08)
            day = campaign_dates[campaign_id]
        if day is None or day < START or day > END:
            continue
        bucket = by_day[day][channel]
        bucket["revenue"] += money(stats.get("conversion_value"))
        if name and name not in bucket["names"]:
            bucket["names"].append(name)
        topic = clean_topic(name)
        if topic and topic not in bucket["topics"]:
            bucket["topics"].append(topic)

    return by_day


def fetch_klaviyo_flows(creds: Credentials):
    headers = klaviyo_headers(creds)
    body = {
        "data": {
            "type": "flow-series-report",
            "attributes": {
                "timeframe": {
                    "start": "2026-03-03T00:00:00-06:00",
                    "end": "2026-04-23T23:59:59-05:00",
                },
                "interval": "daily",
                "conversion_metric_id": KLAVIYO_CONVERSION_METRIC_ID,
                "statistics": ["delivered", "conversion_value"],
                "group_by": ["flow_message_id", "flow_id", "flow_name", "send_channel"],
            },
        }
    }
    report = request_json(
        "https://a.klaviyo.com/api/flow-series-reports/",
        headers=headers,
        method="POST",
        body=body,
    )
    attrs = report["data"]["attributes"]
    dates = [datetime.fromisoformat(x.replace("Z", "+00:00")).date() for x in attrs["date_times"]]
    by_day = defaultdict(lambda: defaultdict(lambda: {"sent": 0.0, "revenue": 0.0}))
    for item in attrs.get("results", []):
        channel = (item.get("groupings", {}).get("send_channel") or "").lower()
        if channel not in {"email", "sms"}:
            continue
        delivered = item["statistics"].get("delivered", [])
        revenue = item["statistics"].get("conversion_value", [])
        for i, day in enumerate(dates):
            if START <= day <= END:
                by_day[day][channel]["sent"] += money(delivered[i] if i < len(delivered) else 0)
                by_day[day][channel]["revenue"] += money(revenue[i] if i < len(revenue) else 0)
    return by_day


def fetch_meta(creds: Credentials):
    out = {d: 0.0 for d in all_dates()}
    url = f"https://graph.facebook.com/v25.0/{creds.meta_account}/insights"
    params = {
        "access_token": creds.meta_token,
        "time_range": json.dumps({"since": START.isoformat(), "until": END.isoformat()}),
        "fields": "date_start,date_stop,spend,action_values,purchase_roas,website_purchase_roas",
        "level": "account",
        "time_increment": "1",
        "limit": "200",
    }
    while url:
        data = request_json(url, params=params if "?" not in url else None)
        params = None
        for row in data.get("data", []):
            day = date.fromisoformat(row["date_start"])
            purchase_value = 0.0
            for action in row.get("action_values", []) or []:
                if action.get("action_type") == "offsite_conversion.fb_pixel_purchase":
                    purchase_value = money(action.get("value"))
                    break
            if not purchase_value:
                for action in row.get("action_values", []) or []:
                    if action.get("action_type") == "purchase":
                        purchase_value = money(action.get("value"))
                        break
            out[day] = purchase_value
        url = data.get("paging", {}).get("next")
    return out


def shopify_day_window(day: date):
    start = datetime(day.year, day.month, day.day, 0, 0, 0, tzinfo=LOCAL_TZ)
    end = start + timedelta(days=1)
    return start.astimezone(timezone.utc).isoformat(), end.astimezone(timezone.utc).isoformat()


def fetch_shopify(creds: Credentials):
    totals = {d: 0.0 for d in all_dates()}
    gql = """
    query($q: String!) {
      shopifyqlQuery(query: $q) {
        tableData {
          columns { name dataType displayName }
          rows
        }
        parseErrors
      }
    }
    """
    shopifyql = (
        "FROM sales SHOW total_sales "
        "WHERE sales_channel = 'Online Store' "
        f"SINCE {START.isoformat()} UNTIL {END.isoformat()} "
        "TIMESERIES day WITH TIMEZONE 'America/Chicago'"
    )
    data = request_json(
        f"https://{creds.shopify_store}/admin/api/2026-01/graphql.json",
        headers={"X-Shopify-Access-Token": creds.shopify_token, "Content-Type": "application/json"},
        method="POST",
        body={"query": gql, "variables": {"q": shopifyql}},
    )
    result = data["data"]["shopifyqlQuery"]
    if result.get("parseErrors"):
        raise RuntimeError(f"ShopifyQL parse errors: {result['parseErrors']}")
    for row in result["tableData"]["rows"]:
        day = date.fromisoformat(row["day"])
        if START <= day <= END:
            totals[day] = money(row.get("total_sales"))
    return totals


def make_workbook(campaigns, flows, meta, shopify):
    wb = Workbook()
    ws = wb.active
    ws.title = "Daily Revenue"

    headers = [
        "Day",
        "Period",
        "Campaigns",
        "Topic",
        "Revenue",
        "Flows sent",
        "Flow revenue",
        "Total email",
        "Campaigns",
        "Topic",
        "Revenue",
        "Flows sent",
        "Flow revenue",
        "Total SMS",
        "Owned channels total",
        "Meta",
        "Other/Organic",
        "Shopify total (Online Store Only)",
    ]
    ws.append(["", "", "Email", "", "", "", "", "", "SMS", "", "", "", "", "", "", "", "", ""])
    ws.append(headers)
    ws.merge_cells(start_row=1, start_column=3, end_row=1, end_column=8)
    ws.merge_cells(start_row=1, start_column=9, end_row=1, end_column=14)

    fill_group = PatternFill("solid", fgColor="D9EAF7")
    fill_header = PatternFill("solid", fgColor="F3F6FA")
    thin = Side(style="thin", color="000000")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for row in ws.iter_rows(min_row=1, max_row=2, min_col=1, max_col=len(headers)):
        for cell in row:
            cell.font = Font(bold=True)
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = border
            cell.fill = fill_group if cell.row == 1 and cell.column in range(3, 15) else fill_header

    for day in all_dates():
        email_campaign = campaigns[day]["email"]
        sms_campaign = campaigns[day]["sms"]
        email_flow = flows[day]["email"]
        sms_flow = flows[day]["sms"]
        email_total = email_campaign["revenue"] + email_flow["revenue"]
        sms_total = sms_campaign["revenue"] + sms_flow["revenue"]
        owned_total = email_total + sms_total
        meta_total = meta.get(day, 0.0)
        shop_total = shopify.get(day, 0.0)
        other = shop_total - owned_total - meta_total
        ws.append(
            [
                f"{day.month}/{day.day}/{day.year}",
                3 if day <= date(2026, 3, 30) else 4,
                "\n".join(email_campaign["names"]),
                "\n".join(email_campaign["topics"]),
                email_campaign["revenue"],
                email_flow["sent"],
                email_flow["revenue"],
                email_total,
                "\n".join(sms_campaign["names"]),
                "\n".join(sms_campaign["topics"]),
                sms_campaign["revenue"],
                sms_flow["sent"],
                sms_flow["revenue"],
                sms_total,
                owned_total,
                meta_total,
                other,
                shop_total,
            ]
        )

    currency_cols = [5, 7, 8, 11, 13, 14, 15, 16, 17, 18]
    number_cols = [6, 12]
    for row in ws.iter_rows(min_row=3, max_row=ws.max_row):
        for cell in row:
            cell.border = border
            cell.alignment = Alignment(vertical="top", wrap_text=True)
        for col in currency_cols:
            ws.cell(row=row[0].row, column=col).number_format = '$#,##0.00;[Red]($#,##0.00)'
        for col in number_cols:
            ws.cell(row=row[0].row, column=col).number_format = '#,##0'

    widths = {
        1: 12,
        2: 8,
        3: 34,
        4: 28,
        5: 12,
        6: 12,
        7: 14,
        8: 14,
        9: 34,
        10: 28,
        11: 12,
        12: 12,
        13: 14,
        14: 14,
        15: 20,
        16: 13,
        17: 15,
        18: 28,
    }
    for col, width in widths.items():
        ws.column_dimensions[get_column_letter(col)].width = width
    ws.freeze_panes = "A3"
    ws.auto_filter.ref = f"A2:R{ws.max_row}"

    notes = wb.create_sheet("Sources & Notes")
    notes.append(["Item", "Detail"])
    notes.append(["Date range", "2026-03-03 through 2026-04-23, America/Chicago"])
    notes.append(
        [
            "Klaviyo campaign revenue",
            "Campaign Values Report, Shopify Placed Order conversion metric, grouped by campaign message and send channel. Campaigns are assigned to the date in the campaign name when present, matching the naming convention in this account.",
        ]
    )
    notes.append(
        [
            "Klaviyo flow revenue/sends",
            "Flow Series Report, daily interval, Shopify Placed Order conversion metric. Flows sent uses delivered.",
        ]
    )
    notes.append(
        [
            "Meta",
            "Meta Ads Insights daily account-level offsite_conversion.fb_pixel_purchase action value.",
        ]
    )
    notes.append(
        [
            "Shopify total",
            "ShopifyQL sales total_sales, filtered to sales_channel = 'Online Store'.",
        ]
    )
    notes.append(
        [
            "Other/Organic",
            "Shopify Online Store total minus Klaviyo email total, Klaviyo SMS total, and Meta. Attribution can overlap across platforms, so this is a directional residual.",
        ]
    )
    notes.append(
        [
            "Direct mail",
            "No PostPilot/API credentials were present in the provided document, so direct mail was not added as an automated column.",
        ]
    )
    for row in notes.iter_rows():
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    notes.column_dimensions["A"].width = 24
    notes.column_dimensions["B"].width = 110
    notes["A1"].font = Font(bold=True)
    notes["B1"].font = Font(bold=True)
    wb.save(OUT)


def main():
    creds = read_credentials()
    print("Fetching Klaviyo campaign performance...", flush=True)
    campaigns = fetch_klaviyo_campaigns(creds)
    print("Fetching Klaviyo flow performance...", flush=True)
    flows = fetch_klaviyo_flows(creds)
    print("Fetching Meta purchase conversion value...", flush=True)
    meta = fetch_meta(creds)
    print("Fetching Shopify online-store totals...", flush=True)
    shopify = fetch_shopify(creds)
    make_workbook(campaigns, flows, meta, shopify)
    print(OUT, flush=True)


if __name__ == "__main__":
    main()
